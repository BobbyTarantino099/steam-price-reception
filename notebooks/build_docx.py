"""Builds entregables/resumen_ejecutivo.docx from the English executive summary.

Mirrors build_pptx.js: content is written directly here rather than parsed from
the Markdown, so the two stay in sync deliberately, not by regex accident.
Run: python notebooks/build_docx.py
"""
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

RUTA_BASE = Path(__file__).resolve().parents[1]
GRAFICOS = RUTA_BASE / 'salidas' / 'graficos'
OUT = RUTA_BASE / 'entregables' / 'resumen_ejecutivo.docx'

NAVY = RGBColor(0x1E, 0x27, 0x61)
ACCENT = RGBColor(0x1F, 0x5F, 0xA8)
DARKTEXT = RGBColor(0x1A, 0x1A, 0x2E)
GRAY = RGBColor(0x6E, 0x7B, 0x91)

doc = Document()
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(11)
doc.styles['Normal'].font.color.rgb = DARKTEXT

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY
        run.font.name = 'Cambria'
    return h

def para(text, *, italic=False, bold=False, size=11, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color or DARKTEXT
    return p

def image(filename, width_in=6.0, caption=None):
    doc.add_picture(str(GRAFICOS / filename), width=Inches(width_in))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        run.italic = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = GRAY

# Title block ---------------------------------------------------------------
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = title.add_run('Price and reception on Steam: where should investment due diligence be prioritised?')
r.bold = True
r.font.size = Pt(20)
r.font.color.rgb = NAVY
r.font.name = 'Cambria'

sub = doc.add_paragraph()
r = sub.add_run('Executive summary — Portfolio case study  ·  July 28, 2026')
r.font.size = Pt(12)
r.font.color.rgb = GRAY

doc.add_paragraph()

# Context ---------------------------------------------------------------
heading('Context', level=2)
para('An investment fund is evaluating entry into the video game sector and needs a data-driven '
     'basis for deciding which genre to prioritise for its investment analysis (due diligence). '
     "This report uses Steam's historical catalogue — the largest PC game market — as a "
     'reference sample of the market.')
p = doc.add_paragraph()
r = p.add_run('Question we answer: '); r.bold = True
p.add_run('which genre and price combinations achieve the best reception from players (measured '
          'as % positive reviews), and which genres are therefore better investment candidates?')

# What we did ---------------------------------------------------------------
heading('What we did', level=2)
para('We analysed 117,430 games published on Steam. To isolate a reliable signal, we focused on '
     'the 8,998 paid games with at least 500 reviews — the minimum volume for the % positive '
     'reviews to be statistically reliable — and grouped them into 10 real video game genres and '
     '4 price bands (from cheapest to most expensive, calculated from the data itself). Prices '
     'were adjusted for inflation so games from different years could be compared on equal '
     'footing.')

# Main finding ---------------------------------------------------------------
heading('Main finding', level=2)
p = doc.add_paragraph()
r = p.add_run('The cheapest price band gets, across all 10 genres without exception, the worst '
              'reception.'); r.bold = True
p.add_run(" The best reception isn't at the highest possible price, but in a mid-to-high range "
          '(approximately $6-$12 adjusted, depending on genre).')
image('01_q1_vs_mejor_franja.png',
      caption='The cheapest band has the worst reception across all 10 Steam genres')
para('The difference is real but moderate: between 2.5 and 5 percentage points of improvement '
     "moving from the cheapest band to each genre's best band. Not a dramatic effect, but a "
     'consistent one across all 10 categories.')
image('02_heatmap_genero_franja.png',
      caption='The best reception clusters at mid-to-high prices, never the cheapest')

# Candidates ---------------------------------------------------------------
heading('Candidates with the best evidence for due diligence', level=2)
para('Of the 10 genres, three combine the best price effect, the largest volume of evidence, and '
     'the highest absolute reception: Adventure, Indie and Casual.')
image('03_ranking_efecto_candidatos.png',
      caption='Adventure, Indie and Casual combine the best effect, volume and absolute reception')
for line in [
    'Adventure shows the largest reception jump between the cheapest band and its best band '
    '(+5 percentage points), with 4,030 games of evidence.',
    'Indie is the genre with the most evidence volume (5,561 games) and a solid improvement '
    '(+3.9 points).',
    'Casual has the highest absolute reception of the 10 genres in its best price band (89.6% '
    'positive reviews), though with lower volume (2,230 games).',
]:
    doc.add_paragraph(line, style='List Bullet')
para("Massively Multiplayer is left off this list: it's the genre with the worst reception "
     'across all price bands and very little evidence (176 games) — not enough data to '
     'confidently recommend or rule it out.')

# Age control ---------------------------------------------------------------
heading('Is this just an effect of cheap games being older?', level=2)
para('No. We repeated the analysis splitting recent games from older ones, and the pattern holds '
     'the same in both groups: the cheapest band remains the worst-reception one, regardless of '
     "the game's age.")
image('04_control_antiguedad.png',
      caption='The pattern holds for recent and older games: it is not an age effect')

# What this doesn't answer ---------------------------------------------------------------
heading("What this analysis doesn't answer", level=2)
for line in [
    "It doesn't prove causality. The pattern shows an association between price and reception, "
    "not that raising a game's price automatically improves its reviews. It may reflect that "
    'studios who charge more also invest more in production quality.',
    'It doesn\'t cover the "long tail." We deliberately excluded games with fewer than 500 '
    'reviews (34% of the catalogue has no reviews at all) to keep statistical reliability. That '
    'tail remains as pending additional analysis.',
    "It doesn't measure reviews over time, only each game's historical cumulative total as of "
    'the data collection date.',
    "This analysis prioritises genres; it doesn't allocate investment budget or replace "
    'studio-or-publisher-specific due diligence.',
]:
    doc.add_paragraph(line, style='List Bullet')

# Next step ---------------------------------------------------------------
heading('Next step', level=2)
para('The formal recommendations phase (Phase 6 of the case) will translate this evidence into a '
     'concrete prioritisation proposal, with its limitations and the additional data recommended '
     'before committing capital.')

doc.add_paragraph()
para('Source: Steam Games Dataset (fronkongames, Kaggle, CC BY 4.0 licence) and the CPI-U '
     'consumer price index (US Bureau of Labor Statistics). Portfolio case study built with '
     'Python/pandas; full methodology and checks available in the technical notebook '
     'notebooks/caso_steam_precio_recepcion.ipynb.',
     italic=True, size=9.5, color=GRAY)

doc.save(str(OUT))
print(f'DOCX generado: {OUT}')
